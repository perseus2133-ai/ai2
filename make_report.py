from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'start', 'bottom', 'end'):
        tag = 'w:' + edge
        element = OxmlElement(tag)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'BFBFBF')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, color='C00000', size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def sub_heading(doc, text, size=13, color='2E4057'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p

def body(doc, text, bullet=False, size=11):
    style = 'List Bullet' if bullet else 'Normal'
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p

def highlight_box(doc, text, bg='FFE0E0', text_color='C00000', size=12):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(text_color)
    doc.add_paragraph()

doc = Document()

# 페이지 여백
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# 제목
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('2028 대입 개편 및 교육 전략 설명회')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('2026년 5월 16일 | 구주이배 교육')
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# ────────────────────────────────────────────
# 1. 2028 수능대비의 핵심
# ────────────────────────────────────────────
heading(doc, '1. 2028 수능대비의 핵심')

sub_heading(doc, '■ 개편(2028~) 수능 과목 및 평가 방식')

tbl = doc.add_table(rows=6, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['과목', '평가 방식']
row0 = tbl.rows[0]
for i, h in enumerate(headers):
    cell = row0.cells[i]
    set_cell_bg(cell, '2E4057')
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

data = [
    ('화법과 언어, 독서와 작문, 문학', '9등급 상대평가'),
    ('대수, 미적분Ⅰ, 확률과 통계', '9등급 상대평가'),
    ('통합사회, 통합과학', '9등급 상대평가 (1학년 과정)'),
    ('영어', '절대평가'),
    ('한국사', '절대평가'),
]
for i, (subject, eval_) in enumerate(data):
    row = tbl.rows[i+1]
    row.cells[0].text = subject
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].text = eval_
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

sub_heading(doc, '■ 핵심 시사점')
body(doc, '변별력이 유지되는 과목: 수학, 국어', bullet=True)
body(doc, '각 대학 입학 전형 시 통합사회·통합과학 반영 비중 축소 예상', bullet=True)
body(doc, '영어, 한국사는 절대평가 → 3등급까지 변별력 없음', bullet=True)
doc.add_paragraph()
highlight_box(doc, '⇒ 수학, 국어가 결정하는 수능', bg='FFE0E0', text_color='C00000')

# ────────────────────────────────────────────
# 2. 새로운 내신체제의 변별력
# ────────────────────────────────────────────
heading(doc, '2. 새로운 내신체제의 변별력')

sub_heading(doc, '■ 내신 체제의 변화')
body(doc, '9등급제 (1등급 4%)  →  5등급제 (1등급 10%)')
body(doc, '등급이 나오는 과목 수 대폭 증가')
doc.add_paragraph()

tbl2 = doc.add_table(rows=3, cols=4)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

header_row = tbl2.rows[0]
for cell in header_row.cells:
    set_cell_bg(cell, '2E4057')
header_row.cells[0].merge(header_row.cells[1])
header_row.cells[0].paragraphs[0].add_run('기존').bold = True
header_row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
header_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
header_row.cells[2].merge(header_row.cells[3])
header_row.cells[2].paragraphs[0].add_run('변경 (2028~)').bold = True
header_row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
header_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

r1_data = [('공통과목', '상대평가(9등급)'), ('일반선택', '상대평가(9등급)')]
r2_data = [('진로선택', '절대평가'), ('전문교과', '절대평가')]
for col, (subj, ev) in enumerate(r1_data):
    tbl2.rows[1].cells[col*2].text = subj
    tbl2.rows[1].cells[col*2+1].text = ev
for col, (subj, ev) in enumerate(r2_data):
    tbl2.rows[2].cells[col*2].text = subj
    tbl2.rows[2].cells[col*2+1].text = ev

doc.add_paragraph()
body(doc, '변경 후: 공통과목·일반선택·진로선택·융합선택·과학고 개설과목 → 모두 상대평가(5등급)', bullet=True)
doc.add_paragraph()

sub_heading(doc, '■ 5등급 변별력 확인 (실증 사례)')
body(doc, '서울 고1 1학기 내신 전 과목 1등급 비율: 1.72% (1,009명)', bullet=True)
body(doc, '남은 과목 중 대부분이 선택과목 → 선택과목 수강생 수 감소로 등급 경쟁 더욱 치열', bullet=True)
doc.add_paragraph()
highlight_box(doc, '⇒ 예상 밖의 역대급 변별력 입증', bg='FFE8CC', text_color='7B3F00')

# ────────────────────────────────────────────
# 3. 2022개정 교육과정과 고교학점제 전면 시행
# ────────────────────────────────────────────
heading(doc, '3. 2022개정 교육과정과 고교학점제 전면 시행')

sub_heading(doc, '■ 고교학점 배당 기준 (주요 교과)')
tbl3 = doc.add_table(rows=7, cols=2)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell in tbl3.rows[0].cells:
    set_cell_bg(cell, '2E4057')
tbl3.rows[0].cells[0].paragraphs[0].add_run('교과(군)').bold = True
tbl3.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
tbl3.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
tbl3.rows[0].cells[1].paragraphs[0].add_run('필수 이수 학점').bold = True
tbl3.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
tbl3.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

subjects_credits = [
    ('국어', '8'), ('수학', '8'), ('영어', '8'),
    ('사회 (역사/도덕 포함)', '6(한국사)/8(통합사회)'),
    ('과학', '10(통합과학, 과학탐구실험)'), ('체육', '10'),
]
for i, (s, c) in enumerate(subjects_credits):
    tbl3.rows[i+1].cells[0].text = s
    tbl3.rows[i+1].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    tbl3.rows[i+1].cells[1].text = c
    tbl3.rows[i+1].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    tbl3.rows[i+1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

sub_heading(doc, '■ 고교 학점제의 현실')
body(doc, '수능과목 필수 수강', bullet=True)
body(doc, '내신 등급에 유리한 과목 개설', bullet=True)
body(doc, '학교 교육 편제 차이가 거의 없음', bullet=True)
doc.add_paragraph()
highlight_box(doc, '"진로에 따른 다양한 과목 선택"은 없다', bg='F0F0F0', text_color='333333')

# ────────────────────────────────────────────
# 4. 무전공 입학
# ────────────────────────────────────────────
heading(doc, '4. 무전공 입학이란?')

sub_heading(doc, '■ 유형 구분')
body(doc, '유형1: 전공을 정하지 않고 모집 후 대학 내 모든 전공 자율 선택 (의전·의료·사범계열 등 제외)', bullet=True)
body(doc, '유형2: 계열 또는 단대 단위 모집 후 전공 자율 선택 (계열단대내 100% 또는 최저정원의 150% 이상)', bullet=True)
doc.add_paragraph()

sub_heading(doc, '■ 대학 인센티브 기준 (전체 정원 내 모집원 기준)')
tbl4 = doc.add_table(rows=3, cols=3)
tbl4.style = 'Table Grid'
for cell in tbl4.rows[0].cells:
    set_cell_bg(cell, '2E4057')
for txt, cell in zip(['학년도', '유형1', '유형1 + 유형2'], tbl4.rows[0].cells):
    p = cell.paragraphs[0]
    p.add_run(txt).bold = True
    p.runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

tbl4.rows[1].cells[0].text = '2025학년도'
tbl4.rows[1].cells[1].text = '5% (단 2026학년도 10% 이상 모집계획 제출 시 인정)'
tbl4.rows[1].cells[2].text = '20% (수도권 사립대) / 25% (국립대)'
tbl4.rows[2].cells[0].text = '2026학년도'
tbl4.rows[2].cells[1].text = '10%'
tbl4.rows[2].cells[2].text = '25% (수도권 사립대) / 30% (국립대)'
doc.add_paragraph()

# ────────────────────────────────────────────
# 5. 2028 대입의 핵심
# ────────────────────────────────────────────
heading(doc, '5. 2028 대입의 핵심')

sub_heading(doc, '■ 핵심 1: 고부터 대입은 시작된다 — 중등과정은 대입 준비과정')

body(doc, '2028 대입 전형의 변화')
doc.add_paragraph()
tbl5 = doc.add_table(rows=3, cols=3)
tbl5.style = 'Table Grid'
for cell in tbl5.rows[0].cells:
    set_cell_bg(cell, '2E4057')
for txt, cell in zip(['전형', '현재', '2028~'], tbl5.rows[0].cells):
    p = cell.paragraphs[0]
    p.add_run(txt).bold = True
    p.runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

tbl5.rows[1].cells[0].text = '수시'
tbl5.rows[1].cells[1].text = '학생부 + 수능'
tbl5.rows[1].cells[2].text = '학생부 + 수능'
tbl5.rows[2].cells[0].text = '정시'
tbl5.rows[2].cells[1].text = '수능'
tbl5.rows[2].cells[2].text = '학생부 + 수능'
doc.add_paragraph()

body(doc, '전형의 구분이 사라짐 → 내신과 수능 모두 잘해야', bullet=True)
body(doc, '패자부활전은 없다', bullet=True)
body(doc, '고등 3년 전 과정이 대입에 반영', bullet=True)
doc.add_paragraph()

sub_heading(doc, '■ 핵심 2: 주요 교과가 결정하는 대입 — 수학·국어 비중 증가')
body(doc, '2028 통합형 수능 반영 방식 (서울대 기준)')
doc.add_paragraph()

tbl6 = doc.add_table(rows=6, cols=2)
tbl6.style = 'Table Grid'
for cell in tbl6.rows[0].cells:
    set_cell_bg(cell, '2E4057')
for txt, cell in zip(['과목', '반영 비율'], tbl6.rows[0].cells):
    p = cell.paragraphs[0]
    p.add_run(txt).bold = True
    p.runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

su_data = [
    ('국어', '100'),
    ('수학', '120 (가장 높음)'),
    ('영어', '등급별 감점 (1등급 0점, 9등급 -14점)'),
    ('한국사', '반영 안 함'),
    ('통합과학·통합사회', '80'),
]
for i, (s, v) in enumerate(su_data):
    tbl6.rows[i+1].cells[0].text = s
    tbl6.rows[i+1].cells[1].text = v
    tbl6.rows[i+1].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    tbl6.rows[i+1].cells[1].paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()
body(doc, '학생부 평가 = 교과 역량 평가 (과목 이수 성도, 학업 성취도, 학습 수행 내용 등)')
doc.add_paragraph()

# ────────────────────────────────────────────
# 6. 최상위 초등의 전략
# ────────────────────────────────────────────
heading(doc, '6. 최상위 초등의 전략')

sub_heading(doc, '■ 진학 전략')
body(doc, '초중등과정에서 고등과정 준비 → 자사고·명문고 전교권 → 메디칼·스카이 명문대 진학', bullet=True)

sub_heading(doc, '■ 학습 전략')
body(doc, '교과 선행 → 고등과정 조기진입 → 내신·수능 1등급 선점', bullet=True)
doc.add_paragraph()

sub_heading(doc, '■ 구체적 전략 구분')
tbl7 = doc.add_table(rows=3, cols=3)
tbl7.style = 'Table Grid'
for cell in tbl7.rows[0].cells:
    set_cell_bg(cell, '2E4057')
for txt, cell in zip(['전략', '경로', '비고'], tbl7.rows[0].cells):
    p = cell.paragraphs[0]
    p.add_run(txt).bold = True
    p.runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

tbl7.rows[1].cells[0].text = '전략1 (메이저)'
tbl7.rows[1].cells[1].text = '초중등→자사고·명문고 전교권→명문대'
tbl7.rows[1].cells[2].text = '명문고 지역 전략 (주류)'
tbl7.rows[2].cells[0].text = '전략2 (마이너)'
tbl7.rows[2].cells[1].text = '초중등→영재고·과고 진학→명문대'
tbl7.rows[2].cells[2].text = '명문고 부재 지역 전략'
doc.add_paragraph()
body(doc, '※ 인문계열 외고 진학은 여전히 유효')
doc.add_paragraph()

# ────────────────────────────────────────────
# 7. 의대목표 낙수효과
# ────────────────────────────────────────────
heading(doc, '7. 의대목표 낙수효과와 최상위권 진로 구조 변화')

sub_heading(doc, '■ 최상위권 진로 구조의 변화')
body(doc, "'이공계 최상위' '연구자·공학자' → 의치약 계열의 소득과 직업 안정성으로 이동", bullet=True)

sub_heading(doc, '■ 의치약 계열 정원 확대')
body(doc, '2022학년도: 약학과 학부 모집 전환', bullet=True)
body(doc, '2025학년도: 의대 모집정원 증가', bullet=True)
body(doc, '의대 모집정원 지속적 증가 (지역의사제 도입)', bullet=True)
doc.add_paragraph()

body(doc, '2027학년도 의치약계열 총 모집인원:')
tbl8 = doc.add_table(rows=2, cols=6)
tbl8.style = 'Table Grid'
for cell in tbl8.rows[0].cells:
    set_cell_bg(cell, '2E4057')
for txt, cell in zip(['의대', '치대', '한의대', '약학대', '수의대', '지역의사'], tbl8.rows[0].cells):
    p = cell.paragraphs[0]
    p.add_run(txt).bold = True
    p.runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)

for txt, cell in zip(['4,485', '630', '725', '1,750', '496', '490'], tbl8.rows[1].cells):
    cell.text = txt
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()
highlight_box(doc, '합계: 총 8,575명 (수능 응시생의 약 1.8%)', bg='FFE0E0', text_color='C00000')

sub_heading(doc, '■ 의대목표 낙수효과')
body(doc, '모든 전형·모든 학과 절대반지 → 서울대 등 상위권대 진학 가능', bullet=True)
doc.add_paragraph()

# ────────────────────────────────────────────
# 8. 영재고 진학의 리스크
# ────────────────────────────────────────────
heading(doc, '8. 영재고 진학의 리스크')

sub_heading(doc, '■ 리스크 요인')
body(doc, '의대 진학 불리: 영재고 출신 의대 진학 제한', bullet=True)
body(doc, '영재고 대비 부담: 올림피아드 준비, 과학 올인, 시험 대비, 연구보고서 작성', bullet=True)
body(doc, '불합격 시 리스크: 자사고·일반고 진학 시 전교권 불확실 (수학 선행 부족, 내신·수능 대비 부족, 국어·영어 절대 열세)', bullet=True)
doc.add_paragraph()

sub_heading(doc, '■ 과고 진학 시')
body(doc, '과고는 카포디지유 전문 → SKY 첨단학과, 계약학과 위주', bullet=True)
body(doc, '2026학년도 입시 기준: 영재고·과고 의대 진학 40%대 급감', bullet=True)
doc.add_paragraph()

# ────────────────────────────────────────────
# 9. 특목고를 대신하는 대안
# ────────────────────────────────────────────
heading(doc, '9. 특목고를 대신하는 효율적인 대안의 등장')

sub_heading(doc, '■ 자사고 대폭 확대')
body(doc, '2010년 이전: 지립형 사립고 (민사고, 광양제철고, 포항제철고 등)', bullet=True)
body(doc, '2010년 이후: 자율형 사립고 대폭 확대', bullet=True)
doc.add_paragraph()

sub_heading(doc, '■ 자사고·명문 일반고의 약진')
body(doc, '의대 진학에서 자사고·명문일반고 조강세', bullet=True)
body(doc, '효율적이고 안정적인 대안으로 등장', bullet=True)
doc.add_paragraph()

body(doc, '2026학년도 서울대 합격자 고교 순위 (상위권): 의대부고·하늘고·황성고·대학임고·단대부고·홍등고·신덕고·세화고·보안고·배재고·진선여고 등')
body(doc, '2025학년도 의대 합격자 (상위권): 삼산고·의론고·단대부고·새화고·중동고·중산고·강서고·화성고·진선여고·심문고·보안고·배재고 등')
doc.add_paragraph()

# ────────────────────────────────────────────
# 10. 상위권 대학 모집인원 구성
# ────────────────────────────────────────────
heading(doc, '10. 상위권 대학 모집인원 구성')

sub_heading(doc, '■ 대학 서열 구조')
body(doc, '① 메이저·수도권 의대', bullet=True)
body(doc, '② 지방의대, 치대', bullet=True)
body(doc, '③ 약대, 한의대, 서울대 자연계, 첨단학과, 인문계 상위권 학과', bullet=True)
body(doc, '④ 연고대 중하위권 학과', bullet=True)
body(doc, '⑤ 서성한이, 중경외시', bullet=True)
doc.add_paragraph()
highlight_box(doc, '자연계 약 11,000명 : 인문계 약 1,100명 (약 10:1)', bg='E8F4FF', text_color='1A3A5C')

# ────────────────────────────────────────────
# 11. H고 과목선택 및 주요 과목 비중
# ────────────────────────────────────────────
heading(doc, '11. H고 과목선택 → 주요 과목 비중 증가')

sub_heading(doc, '■ 과목별 이수 목록 (H고 기준)')
tbl9 = doc.add_table(rows=5, cols=2)
tbl9.style = 'Table Grid'
for cell in tbl9.rows[0].cells:
    set_cell_bg(cell, '2E4057')
for txt, cell in zip(['교과', '이수 과목'], tbl9.rows[0].cells):
    p = cell.paragraphs[0]
    p.add_run(txt).bold = True
    p.runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

subj_data = [
    ('국어 (8과목)', '공통국어1·2 / 문학 / 독서와 작문 / 화법과 언어 / 언어생활과 탐구 / 매체 의사와 소통 / 독서토론과 글쓰기'),
    ('영어 (5과목)', '공통영어1·2 / 영어1·2 / 영어 독해와 작문'),
    ('수학 (10과목)', '공통수학1·2 / 대수 / 미적분Ⅰ / 확률과 통계 / 기하(경제수학) / 미적분Ⅱ / 고급 미적분 / 인공지능수학 / 수학과제탐구'),
    ('과학 (8.5과목)', '통합과학1·2 / 물리학 / 화학 / 생명과학 / 역학과 에너지 / 화학반응의 세계 / 물질과 에너지 / 세포와 물질대사 / 생물의 유전 / 고급 화학'),
]
for i, (subj, content) in enumerate(subj_data):
    tbl9.rows[i+1].cells[0].text = subj
    tbl9.rows[i+1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    tbl9.rows[i+1].cells[1].text = content
    tbl9.rows[i+1].cells[1].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()
highlight_box(doc, '"내신도 수능도" 수학 비중이 절대적', bg='FFE0E0', text_color='C00000')

# ────────────────────────────────────────────
# 12. 고교 진학 후 수학 이수 전략
# ────────────────────────────────────────────
heading(doc, '12. 고교 진학 후 수학 이수 전략')

sub_heading(doc, '■ 핵심 원칙: 내신 상위권 선행이 필수인 구조')
body(doc, '무거운 과목(수능과목): 수능과목 완성, 내신문제=수능문제, 겨울방학 수능 준비', bullet=True)
body(doc, '심화 과목: 심화 과목 이수 경쟁, 대학 권장 과목 이수', bullet=True)
body(doc, '미적분Ⅱ, 기하 선행 필수', bullet=True)
doc.add_paragraph()

sub_heading(doc, '■ H고 / M고 수학 이수 예시')
tbl10 = doc.add_table(rows=3, cols=7)
tbl10.style = 'Table Grid'
for cell in tbl10.rows[0].cells:
    set_cell_bg(cell, '2E4057')
for txt, cell in zip(['학교', '1-1', '1-2', '2-1', '2-2', '3-1', '3-2'], tbl10.rows[0].cells):
    p = cell.paragraphs[0]
    p.add_run(txt).bold = True
    p.runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)

hgo = ['H고', '공통수학1', '공통수학2', '대수', '미적분Ⅰ', '미적분Ⅱ/확률과통계/기하 등', '인공지능수학/고급미적분/수학과제탐구 등']
mgo = ['M고', '공통수학1', '공통수학2', '대수', '미적분Ⅰ/기하', '확률과통계/미적분Ⅱ/경제수학', '수학과제탐구/실용통계 등']
for i, row_data in enumerate([hgo, mgo]):
    for j, txt in enumerate(row_data):
        tbl10.rows[i+1].cells[j].text = txt
        tbl10.rows[i+1].cells[j].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

# ────────────────────────────────────────────
# 13. 구주이배 수준별 커리큘럼
# ────────────────────────────────────────────
heading(doc, '13. 구주이배 수준별 수학 커리큘럼')

sub_heading(doc, '■ 고등학교 진학 전 수준별 커리큘럼 (Z·M·G 트랙)')
tbl11 = doc.add_table(rows=8, cols=7)
tbl11.style = 'Table Grid'
for cell in tbl11.rows[0].cells:
    set_cell_bg(cell, '2E4057')
for txt, cell in zip(['레벨', '공통수학', '내수', '미적분Ⅰ', '확률과통계', '기하', '미적분Ⅱ'], tbl11.rows[0].cells):
    p = cell.paragraphs[0]
    p.add_run(txt).bold = True
    p.runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)

levels = [
    ['Z', '내신실전', '수능실전', '수능실전', '수능실전', '기본', '기본'],
    ['M', '내신실전', '수능실전', '수능실전', '심화', '기본', ''],
    ['M', '내신실전', '심화', '심화', '기본', '', ''],
    ['G', '내신실전', '심화', '기본', '', '', ''],
    ['G', '내신실전', '기본', '기본', '', '', ''],
    ['G', '내신실전', '기본', '', '', '', ''],
    ['G', '내신실전', '', '', '', '', ''],
]
for i, row_data in enumerate(levels):
    for j, txt in enumerate(row_data):
        tbl11.rows[i+1].cells[j].text = txt
        if len(tbl11.rows[i+1].cells[j].paragraphs[0].runs) > 0:
            tbl11.rows[i+1].cells[j].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()
body(doc, '선행 필수 과목: 상위권 결정 과목 = 고등 주요 내신 결정 = 수능과목', bullet=True)
body(doc, '최신 선행 전략: 대수·미적분·확통·심화 중심', bullet=True)
doc.add_paragraph()

sub_heading(doc, '■ 수준별 수학 커리큘럼 (초6~중3 시기별)')
tbl12 = doc.add_table(rows=6, cols=3)
tbl12.style = 'Table Grid'
for cell in tbl12.rows[0].cells:
    set_cell_bg(cell, '2E4057')
for txt, cell in zip(['레벨', '시작 시점', '고등과정 준비 기간'], tbl12.rows[0].cells):
    p = cell.paragraphs[0]
    p.add_run(txt).bold = True
    p.runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

track_data = [
    ('Z', '초6에 공통수학1 시작 (중1부터 고등과정)', '고등과정 4년'),
    ('M', '중등과정 후 공수1 시작', '고등과정 3~3.5년'),
    ('M', '중등과정 → 공수1 (중1)', '고등과정 3년'),
    ('G', '중등과정 → 공수1 (중2)', '고등과정 2.5년'),
    ('G', '중등과정 → 공수1 (중3)', '고등과정 2년'),
]
for i, (lv, start, period) in enumerate(track_data):
    tbl12.rows[i+1].cells[0].text = lv
    tbl12.rows[i+1].cells[1].text = start
    tbl12.rows[i+1].cells[2].text = period
    for cell in tbl12.rows[i+1].cells:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()
highlight_box(doc, '핵심: 고등과정 시작이 선행 결과를 결정한다\n상위 20개 대학 목표 시 고등선행 최소 2년 확보 필요', bg='E8FFE8', text_color='1A5C1A')

# ────────────────────────────────────────────
# 14. 수학 과정별 학습량과 난이도
# ────────────────────────────────────────────
heading(doc, '14. 수학 과정별 학습량과 난이도')

sub_heading(doc, '■ 단계적 학습 구조 (학습량 ↑, 난이도 ↑)')
tbl13 = doc.add_table(rows=5, cols=2)
tbl13.style = 'Table Grid'
for cell in tbl13.rows[0].cells:
    set_cell_bg(cell, '2E4057')
for txt, cell in zip(['단계', '학습 내용'], tbl13.rows[0].cells):
    p = cell.paragraphs[0]
    p.add_run(txt).bold = True
    p.runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

stages = [
    ('1단계', '초등과정 기초'),
    ('2단계', '중등과정 기초'),
    ('3단계', '고등과정 개념 및 유형'),
    ('4단계', '고등내신 1등급 문제해결 학습'),
    ('5단계 (최상위)', '수능 1등급 문제해결 학습'),
]
for i, (stage, content) in enumerate(stages):
    tbl13.rows[i].cells[0].text = stage
    tbl13.rows[i].cells[1].text = content

doc.add_paragraph()

# 마무리
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_end = p_end.add_run('─ 끝 ─')
r_end.font.size = Pt(11)
r_end.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save('/home/user/ai2/2028대입_교육전략_보고자료.docx')
print("Word 저장 완료")
